"""
Word 报告生成器 —— 使用 python-docx 生成 .docx 文件
"""

import os
import uuid
from datetime import datetime
from typing import List, Dict, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "outputs")


class ReportGenerator:
    """Word 报告生成器"""

    def generate(self, analysis_id: str, markdown: str, 
                 cases: List[Dict]) -> Tuple[str, str]:
        """
        生成 Word 文档
        
        返回：(file_id, filename)
        """
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        now = datetime.now()
        filename = f"legal_analysis_{now.strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(OUTPUT_DIR, filename)

        doc = Document()

        # ── 页面设置 ──
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

        # ── 标题 ──
        title = doc.add_heading("法律分析报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(f"生成时间：{now.strftime('%Y年%m月%d日 %H:%M:%S')}")
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # 空行

        # ── 按段落处理 Markdown ──
        lines = markdown.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("# ") and not line.startswith("## "):
                # 一级标题已处理
                continue
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=3)
            elif line.startswith("- "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.runs[0].font.italic = True if p.runs else True
            elif line == "---":
                doc.add_paragraph("_" * 50)
            else:
                # 普通段落 - 处理粗体标记
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

        # ── 案例附录 ──
        if cases:
            doc.add_page_break()
            doc.add_heading("附录：引用案例清单", level=1)
            for i, case in enumerate(cases, 1):
                mock_tag = " [模拟案例]" if case.get("is_mock") else ""
                doc.add_heading(f"案例 {i}{mock_tag}", level=2)
                
                fields = [
                    ("标题", case.get("title", "")),
                    ("案号", case.get("case_no", "")),
                    ("法院", case.get("court", "")),
                    ("裁判日期", case.get("judgment_date", "")),
                    ("争议焦点", case.get("issue", "")),
                    ("裁判观点", case.get("holding", "")),
                    ("链接", case.get("url", "")),
                    ("数据源", case.get("source_name", "")),
                ]
                for label, value in fields:
                    if value:
                        p = doc.add_paragraph()
                        run_label = p.add_run(f"{label}：")
                        run_label.bold = True
                        p.add_run(value)

        # ── 保存 ──
        doc.save(file_path)

        # 生成 file_id（取文件名不含扩展名部分 + 短 hash）
        file_id = f"{now.strftime('%Y%m%d%H%M%S')}_{analysis_id}"

        return file_id, filename

    @staticmethod
    def _add_formatted_text(paragraph, text: str):
        """简单处理 **粗体** 文本"""
        import re
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    @staticmethod
    def get_file_path(filename: str) -> str:
        return os.path.join(OUTPUT_DIR, filename)

    @staticmethod
    def file_exists(filename: str) -> bool:
        return os.path.exists(os.path.join(OUTPUT_DIR, filename))


def get_report_generator() -> ReportGenerator:
    return ReportGenerator()
